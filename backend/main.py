import os
import uuid
import logging
import subprocess
import shutil
import imgkit
from datetime import datetime
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from supabase import create_client, Client
from dotenv import load_dotenv

# Библиотеки для работы с Word
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Загрузка переменных окружения
load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# Разрешаем все источники, чтобы гарантированно убрать ошибку
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Инициализация Supabase
SUPABASE_URL = os.getenv("VITE_SUPABASE_URL")
SUPABASE_KEY = os.getenv("VITE_SUPABASE_ANON_KEY")
# ВАЖНО: Для админ-функций (confirm_payment) лучше использовать SERVICE_ROLE_KEY в .env
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# --- МОДЕЛИ ДАННЫХ ---

class ReportRequest(BaseModel):
    repoUrl: str
    fullName: str
    group: str
    workType: str
    workNumber: str
    variant: str
    userId: str
    instruction: str

class SBPRequest(BaseModel):
    userId: str
    amount: int

# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---

def set_font_times(run, size=14, bold=False):
    """Установка шрифта Times New Roman согласно ГОСТ МГТУ"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold

def parse_instruction(text):
    """Разбор текста на Тему, Цель и Задачи"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    theme = lines[0] if lines else "Отчет по работе"
    goal = "Выполнить работу в соответствии с заданием."
    tasks = text
    
    for line in lines:
        if line.lower().startswith("цель:"):
            goal = line.split(":", 1)[1].strip()
            
    return theme, goal, tasks

def process_github_content(repo_url, temp_dir):
    """Клонирование репозитория и извлечение визуализаций"""
    images = []
    try:
        # Клонируем репо
        subprocess.run(["git", "clone", repo_url, temp_dir], check=True, timeout=60, capture_output=True)
        
        for root, dirs, files in os.walk(temp_dir):
            if any(x in root for x in ['venv', '.git', '__pycache__']):
                continue
            for file in files:
                full_path = os.path.join(root, file)
                # Ищем картинки
                if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                    images.append(full_path)
                # Ищем HTML графики и конвертируем в PNG
                elif file.lower().endswith('.html'):
                    try:
                        img_path = full_path.replace('.html', '_snap.png')
                        imgkit.from_file(full_path, img_path, options={'quiet': ''})
                        images.append(img_path)
                    except Exception as e:
                        logger.warning(f"Ошибка конвертации HTML {file}: {e}")
    except Exception as e:
        logger.error(f"Ошибка работы с GitHub: {e}")
    return images

# --- ЭНДПОИНТЫ ОПЛАТЫ (СБП) ---

@app.post("/create-sbp-request")
async def create_sbp(request: SBPRequest):
    """Создает запись о намерении оплаты в БД"""
    try:
        res = supabase.table("payment_requests").insert({
            "user_id": request.userId,
            "amount": request.amount,
            "status": "pending"
        }).execute()
        return {"status": "success", "paymentId": res.data[0]['id']}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/admin/confirm-payment/{payment_id}")
async def confirm_payment(payment_id: str):
    """Ручное подтверждение оплаты администратором"""
    try:
        # Находим заявку
        pay_res = supabase.table("payment_requests").select("*").eq("id", payment_id).single().execute()
        if not pay_res.data:
            raise HTTPException(status_code=404, detail="Заявка не найдена")
        
        user_id = pay_res.data['user_id']
        
        # Обновляем профиль пользователя до PRO
        supabase.table("user_profiles").update({
            "plan": "pro",
            "reports_generated": 0 # Сбрасываем счетчик при покупке
        }).eq("id", user_id).execute()
        
        # Закрываем заявку
        supabase.table("payment_requests").update({"status": "completed"}).eq("id", payment_id).execute()
        
        return {"status": "success", "message": f"User {user_id} upgraded to PRO"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- ОСНОВНОЙ ЭНДПОИНТ ГЕНЕРАЦИИ ---

@app.post("/generate-report")
async def generate_report(request: ReportRequest):
    # 1. Проверка лимитов пользователя
    user_res = supabase.table("user_profiles").select("*").eq("id", request.userId).execute()
    if not user_res.data:
        # Создаем профиль по умолчанию, если его нет
        supabase.table("user_profiles").insert({"id": request.userId, "plan": "free"}).execute()
        user_data = {"plan": "free", "reports_generated": 0}
    else:
        user_data = user_res.data[0]

    if user_data['plan'] == 'free' and user_data['reports_generated'] >= 3:
        raise HTTPException(status_code=403, detail="Лимит бесплатных отчетов (3) исчерпан. Пожалуйста, приобретите PRO.")

    repo_dir = f"temp/repo_{uuid.uuid4().hex}"
    local_docx_path = f"temp/Report_{uuid.uuid4().hex[:6]}.docx"
    
    try:
        # 2. Обработка данных
        images = process_github_content(request.repoUrl, repo_dir)
        theme, goal, tasks = parse_instruction(request.instruction)
        
        doc = Document()
        
        # Настройка полей ГОСТ МГТУ: левое 30мм, правое 10мм
        section = doc.sections[0]
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.39)

        # --- ТИТУЛЬНЫЙ ЛИСТ ---
        p_header = doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        univ_header = (
            "Министерство науки и высшего образования Российской Федерации\n"
            "Федеральное государственное автономное образовательное учреждение\n"
            "высшего образования\n"
            "«Московский государственный технический университет\n"
            "имени Н.Э. Баумана\n"
            "(национальный исследовательский университет)»\n"
            "(МГТУ им. Н.Э. Баумана)\n"
        )
        set_font_times(p_header.add_run(univ_header), 10)

        p_fac = doc.add_paragraph()
        p_fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fac = p_fac.add_run("\nФАКУЛЬТЕТ ИНФОРМАТИКА И СИСТЕМЫ УПРАВЛЕНИЯ\n"
                                "КАФЕДРА СИСТЕМЫ ОБРАБОТКИ ИНФОРМАЦИИ И УПРАВЛЕНИЯ\n")
        set_font_times(run_fac, 12, bold=True)

        for _ in range(4): doc.add_paragraph()
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_times(p_title.add_run(f"{request.workType.upper()} №{request.workNumber}\nНА ТЕМУ:"), 16)
        
        p_theme = doc.add_paragraph()
        p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_times(p_theme.add_run(f"«{theme}»"), 14, bold=True)

        for _ in range(5): doc.add_paragraph()
        
        # Таблица подписей
        table = doc.add_table(rows=2, cols=3)
        table.columns[0].width = Inches(2.5)
        
        # Студент
        c1 = table.rows[0].cells
        set_font_times(c1[0].paragraphs[0].add_run(f"Студент {request.group}"), 12)
        set_font_times(c1[2].paragraphs[0].add_run(request.fullName), 12)
        
        # Преподаватель
        c2 = table.rows[1].cells
        set_font_times(c2[0].paragraphs[0].add_run("Преподаватель"), 12)
        set_font_times(c2[2].paragraphs[0].add_run("М.И. Колосов"), 12)

        # Год
        for _ in range(4): doc.add_paragraph()
        p_year = doc.add_paragraph(f"{datetime.now().year} г.")
        p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_times(p_year.runs[0], 12)

        doc.add_page_break()

        # --- СОДЕРЖАТЕЛЬНАЯ ЧАСТЬ ---
        doc.add_heading('1. ЦЕЛЬ И ЗАДАЧИ РАБОТЫ', level=1)
        p_goal = doc.add_paragraph()
        set_font_times(p_goal.add_run(f"Цель работы: {goal}"), 14)
        
        doc.add_heading('2. ФОРМУЛИРОВКА ЗАДАНИЯ ПО ВАРИАНТУ', level=1)
        p_tasks = doc.add_paragraph()
        p_tasks.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_tasks.paragraph_format.first_line_indent = Inches(0.5)
        set_font_times(p_tasks.add_run(f"Вариант {request.variant} включает выполнение следующего задания:\n{tasks}"), 14)

        # --- РЕЗУЛЬТАТЫ (ГРАФИКИ) ---
        if images:
            doc.add_heading('3. РЕЗУЛЬТАТЫ ВЫПОЛНЕНИЯ', level=1)
            for idx, img_path in enumerate(images[:5]): # Не более 5 картинок
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                    p_img = doc.add_paragraph(f"Рисунок {idx+1} — Графический результат из репозитория")
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font_times(p_img.runs[0], 12)
                except Exception as img_err:
                    logger.warning(f"Не удалось вставить картинку {img_path}: {img_err}")

        # Сохранение локально
        doc.save(local_docx_path)

        # 3. Загрузка в Supabase Storage и обновление статистики
        storage_filename = f"{request.userId}/{os.path.basename(local_docx_path)}"
        with open(local_docx_path, "rb") as f:
            supabase.storage.from_("reports").upload(storage_filename, f)

        public_url = supabase.storage.from_("reports").get_public_url(storage_filename)
        
        # Инкремент счетчика отчетов
        supabase.table("user_profiles").update({
            "reports_generated": user_data['reports_generated'] + 1
        }).eq("id", request.userId).execute()

        # Запись в историю отчетов
        supabase.table("reports").insert({
            "user_id": request.userId,
            "topic": f"{request.workType} №{request.workNumber}",
            "file_url": public_url,
            "github_url": request.repoUrl
        }).execute()

        return {"status": "success", "downloadUrl": public_url}

    except Exception as e:
        logger.error(f"Глобальная ошибка генерации: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        # Очистка временных файлов
        if os.path.exists(repo_dir):
            shutil.rmtree(repo_dir, ignore_errors=True)
        if os.path.exists(local_docx_path):
            os.remove(local_docx_path)

if __name__ == "__main__":
    import uvicorn
    # Запуск сервера на порту 8000
    uvicorn.run(app, host="0.0.0.0", port=8000)